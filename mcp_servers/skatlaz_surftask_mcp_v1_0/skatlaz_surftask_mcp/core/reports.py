from __future__ import annotations
from pathlib import Path
from docx import Document
import pandas as pd

def make_docx_report(page: dict, path: str | Path) -> str:
    path = Path(path); path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(page.get("title") or "Relatório Skatlaz SurfTask", 0)
    doc.add_paragraph(f"URL: {page.get('url','')}")
    doc.add_heading("Descrição", level=1)
    doc.add_paragraph(page.get("description") or "")
    doc.add_heading("Resumo", level=1)
    doc.add_paragraph(page.get("summary") or "")
    doc.add_heading("Headings", level=1)
    for h in page.get("headings", [])[:20]:
        doc.add_paragraph(h, style="List Bullet")
    doc.add_heading("Metatags", level=1)
    for k, v in list(page.get("metatags", {}).items())[:50]:
        doc.add_paragraph(f"{k}: {v}")
    doc.save(path)
    return str(path)

def make_xlsx_report(pages: list[dict], path: str | Path) -> str:
    path = Path(path); path.parent.mkdir(parents=True, exist_ok=True)
    rows = []
    for p in pages:
        rows.append({"title": p.get("title"), "url": p.get("url"), "description": p.get("description"), "summary": p.get("summary")})
    pd.DataFrame(rows).to_excel(path, index=False)
    return str(path)
