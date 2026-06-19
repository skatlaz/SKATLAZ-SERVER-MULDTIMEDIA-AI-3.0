from pathlib import Path
from fpdf import FPDF
from docx import Document


def text_to_pdf(text: str, output_path: str, title: str = 'Skatlaz OpenStudio') -> str:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', 'B', 16)
    pdf.multi_cell(0, 10, title)
    pdf.set_font('Arial', '', 12)
    pdf.multi_cell(0, 8, text.encode('latin-1', errors='replace').decode('latin-1'))
    pdf.output(output_path)
    return output_path


def text_to_docx(text: str, output_path: str, title: str = 'Skatlaz OpenStudio') -> str:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    doc.save(output_path)
    return output_path
